#!/usr/bin/env python3
"""Extract reviewable text from common PSP offer files.

This is the adapter layer before the deterministic OfferPSP normalizer. It never
publishes an offer and never guesses missing commercial terms. The original file
remains the audit source; extracted text is only an intermediate draft input.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import mimetypes
from pathlib import Path
from typing import Iterable


TEXT_SUFFIXES = {
    ".txt", ".md", ".csv", ".tsv", ".json", ".xml", ".html", ".htm",
}


def compact_lines(lines: Iterable[str]) -> str:
    cleaned = [line.replace("\u00a0", " ").rstrip() for line in lines]
    return "\n".join(cleaned).strip()


def extract_delimited(path: Path, delimiter: str) -> str:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = csv.reader(handle, delimiter=delimiter)
        return compact_lines(" | ".join(cell.strip() for cell in row) for row in rows)


def extract_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as error:
        raise RuntimeError("XLSX extraction requires openpyxl.") from error

    workbook = load_workbook(path, read_only=True, data_only=True)
    sections: list[str] = []
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(value).strip() for value in row if value not in (None, "")]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            sections.append(f"Sheet: {sheet.title}\n" + "\n".join(rows))
    return compact_lines(sections)


def extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise RuntimeError("PDF extraction requires pypdf.") from error

    reader = PdfReader(path)
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"Page {index}\n{text}")
    if not pages:
        raise RuntimeError("The PDF contains no extractable text. OCR/manual review is required.")
    return compact_lines(pages)


def extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError("DOCX extraction requires python-docx.") from error

    document = Document(path)
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return compact_lines(lines)


def extract_source(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix == ".csv":
        return extract_delimited(path, ","), "csv"
    if suffix == ".tsv":
        return extract_delimited(path, "\t"), "tsv"
    if suffix in TEXT_SUFFIXES:
        return path.read_text(encoding="utf-8-sig"), suffix.lstrip(".") or "text"
    if suffix == ".xlsx":
        return extract_xlsx(path), "xlsx"
    if suffix == ".pdf":
        return extract_pdf(path), "pdf"
    if suffix == ".docx":
        return extract_docx(path), "docx"
    raise RuntimeError(f"Unsupported source format: {suffix or 'unknown'}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--metadata")
    args = parser.parse_args()

    source = Path(args.source).expanduser().resolve()
    output = Path(args.output).expanduser().resolve()
    if not source.is_file():
        raise RuntimeError(f"Source file does not exist: {source}")

    extracted, source_format = extract_source(source)
    extracted = compact_lines(extracted.splitlines())
    if not extracted:
        raise RuntimeError("The source produced no reviewable text.")

    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(extracted + "\n", encoding="utf-8")
    output.chmod(0o600)

    metadata_path = Path(args.metadata).expanduser().resolve() if args.metadata else output.with_suffix(output.suffix + ".meta.json")
    metadata = {
        "ingestion_standard": "offerpsp-universal-source-v1",
        "source_file_name": source.name,
        "source_format": source_format,
        "source_mime_type": mimetypes.guess_type(source.name)[0],
        "source_size_bytes": source.stat().st_size,
        "original_source_sha256": hashlib.sha256(source.read_bytes()).hexdigest(),
        "extraction_method": f"offerpsp-file-adapter:{source_format}",
        "extractor_version": "offerpsp-source-extractor-v1",
        "extracted_text_sha256": hashlib.sha256(extracted.encode("utf-8")).hexdigest(),
        "publication_allowed": False,
        "requires_staff_review": True,
    }
    metadata_path.parent.mkdir(parents=True, exist_ok=True)
    metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    metadata_path.chmod(0o600)
    print(json.dumps({"output": str(output), "metadata": str(metadata_path), "source_format": source_format}, ensure_ascii=False))


if __name__ == "__main__":
    main()
